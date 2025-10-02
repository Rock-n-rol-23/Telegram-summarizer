import os
import tempfile
import logging
import aiofiles
import aiohttp
from typing import Dict, Any, Optional
import chardet
import re

# Импорты для работы с разными форматами файлов
try:
    import PyPDF2
    import pdfplumber
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

try:
    from docx import Document
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

try:
    import mammoth
    HAS_DOC_SUPPORT = True
except ImportError:
    HAS_DOC_SUPPORT = False

# Импорты для новых модулей
try:
    from content_extraction.pdf_ocr import extract_text_from_pdf as ocr_pdf_extract
    from content_extraction.pptx_extractor import extract_text_from_pptx
    HAS_ENHANCED_PDF_SUPPORT = True
    HAS_PPTX_SUPPORT = True
except ImportError as e:
    HAS_ENHANCED_PDF_SUPPORT = False
    HAS_PPTX_SUPPORT = False
    logger.warning(f"Улучшенная поддержка PDF/PPTX недоступна: {e}")

# Импорты для книжных форматов
try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    HAS_EPUB_SUPPORT = True
except ImportError:
    HAS_EPUB_SUPPORT = False

try:
    import xml.etree.ElementTree as ET
    HAS_FB2_SUPPORT = True
except ImportError:
    HAS_FB2_SUPPORT = False

logger = logging.getLogger(__name__)

class FileProcessor:
    """Класс для обработки файлов и извлечения текста"""
    
    def __init__(self):
        # Расширяем поддерживаемые форматы
        base_formats = ['.pdf', '.docx', '.doc', '.txt']
        if HAS_PPTX_SUPPORT:
            base_formats.append('.pptx')
        if HAS_ENHANCED_PDF_SUPPORT:
            base_formats.extend(['.png', '.jpg', '.jpeg'])  # OCR для изображений
        if HAS_EPUB_SUPPORT:
            base_formats.append('.epub')
        if HAS_FB2_SUPPORT:
            base_formats.append('.fb2')

        self.supported_extensions = base_formats
        self.max_file_size = 20 * 1024 * 1024  # 20MB - лимит Telegram
        
    async def download_telegram_file(self, file_info: Dict[str, Any], file_name: str, file_size: int) -> Dict[str, Any]:
        """Скачивает файл от Telegram бота"""
        try:
            # Проверяем размер файла
            if file_size > self.max_file_size:
                return {
                    'success': False,
                    'error': 'Файл слишком большой (максимум 20MB)'
                }
            
            # Проверяем расширение файла
            file_extension = os.path.splitext(file_name.lower())[1]
            if file_extension not in self.supported_extensions:
                return {
                    'success': False,
                    'error': f'Неподдерживаемый формат файла. Поддерживаются: {", ".join(self.supported_extensions)}'
                }
            
            # Создаем временную директорию
            temp_dir = tempfile.mkdtemp()
            local_file_path = os.path.join(temp_dir, file_name)
            
            # Скачиваем файл
            async with aiohttp.ClientSession() as session:
                async with session.get(file_info['file_path']) as response:
                    if response.status == 200:
                        async with aiofiles.open(local_file_path, 'wb') as f:
                            async for chunk in response.content.iter_chunked(8192):
                                await f.write(chunk)
                    else:
                        return {
                            'success': False,
                            'error': f'Ошибка скачивания файла: HTTP {response.status}'
                        }
            
            return {
                'success': True,
                'file_path': local_file_path,
                'file_name': file_name,
                'file_size': file_size,
                'file_extension': file_extension,
                'temp_dir': temp_dir
            }
            
        except Exception as e:
            logger.error(f"Ошибка при скачивании файла: {e}")
            return {
                'success': False,
                'error': f'Ошибка при скачивании файла: {str(e)}'
            }
    
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из PDF файла с поддержкой OCR"""
        # Приоритет новому экстрактору с OCR
        if HAS_ENHANCED_PDF_SUPPORT:
            logger.info("📄 Используется улучшенный PDF экстрактор с OCR")
            return ocr_pdf_extract(
                file_path,
                ocr_langs=os.getenv("OCR_LANGS", "rus+eng"),
                dpi=int(os.getenv("PDF_OCR_DPI", "200")),
                max_pages_ocr=int(os.getenv("MAX_PAGES_OCR", "50"))
            )
        
        # Fallback на старый метод
        if not HAS_PDF_SUPPORT:
            return {
                'success': False,
                'error': 'PDF библиотеки не установлены'
            }
        
        try:
            logger.info(f"📄 Начинаю извлечение текста из PDF: {file_path}")
            text = ""
            
            # Сначала пробуем pdfplumber (лучше для сложных PDF)
            try:
                logger.info("📄 Пробую pdfplumber...")
                with pdfplumber.open(file_path) as pdf:
                    logger.info(f"📄 PDF открыт, страниц: {len(pdf.pages)}")
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                            logger.info(f"📄 Обработана страница {i+1}, символов: {len(page_text)}")
                            
                if text.strip():
                    logger.info(f"📄 pdfplumber успешно, извлечено {len(text)} символов")
                    return {
                        'success': True,
                        'text': text.strip(),
                        'method': 'pdfplumber'
                    }
                else:
                    logger.warning("📄 pdfplumber не извлек текст")
            except Exception as e:
                logger.warning(f"📄 pdfplumber не сработал: {e}")
            
            # Fallback на PyPDF2
            try:
                logger.info("📄 Пробую PyPDF2...")
                text = ""  # Сбрасываем текст для PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    logger.info(f"📄 PyPDF2 PDF открыт, страниц: {len(reader.pages)}")
                    
                    # Проверяем, зашифрован ли PDF
                    if reader.is_encrypted:
                        logger.warning("📄 PDF файл защищен паролем")
                        return {
                            'success': False,
                            'error': 'PDF файл защищен паролем'
                        }
                    
                    for i, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n\n"
                                logger.info(f"📄 PyPDF2 обработана страница {i+1}, символов: {len(page_text)}")
                        except Exception as page_error:
                            logger.warning(f"📄 Ошибка обработки страницы {i+1}: {page_error}")
                            continue
                        
                if text.strip():
                    logger.info(f"📄 PyPDF2 успешно, извлечено {len(text)} символов")
                    return {
                        'success': True,
                        'text': text.strip(),
                        'method': 'PyPDF2'
                    }
                else:
                    logger.warning("📄 PyPDF2 не извлек текст")
                    return {
                        'success': False,
                        'error': 'PDF не содержит извлекаемого текста (возможно, только изображения)'
                    }
                    
            except Exception as e:
                logger.error(f"📄 Ошибка PyPDF2: {e}")
                return {
                    'success': False,
                    'error': f'Ошибка чтения PDF: {str(e)}'
                }
                
        except Exception as e:
            logger.error(f"📄 Критическая ошибка обработки PDF: {e}")
            return {
                'success': False,
                'error': f'Ошибка обработки PDF: {str(e)}'
            }
    
    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из DOCX файла"""
        if not HAS_DOCX_SUPPORT:
            return {
                'success': False,
                'error': 'DOCX библиотека не установлена'
            }
        
        try:
            if not HAS_DOCX_SUPPORT:
                return {
                    'success': False,
                    'error': 'DOCX библиотеки не установлены'
                }
            doc = Document(file_path)
            text = ""
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                return {
                    'success': True,
                    'text': text.strip(),
                    'method': 'python-docx'
                }
            else:
                return {
                    'success': False,
                    'error': 'DOCX файл не содержит текста'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'Ошибка чтения DOCX: {str(e)}'
            }
    
    def extract_text_from_doc(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из DOC файла (старый формат Word)"""
        if not HAS_DOC_SUPPORT:
            return {
                'success': False,
                'error': 'DOC библиотека не установлена'
            }
        
        try:
            if not HAS_DOC_SUPPORT:
                return {
                    'success': False,
                    'error': 'DOC библиотеки не установлены'
                }
            # Пробуем mammoth для DOC файлов
            with open(file_path, "rb") as doc_file:
                result = mammoth.convert_to_html(doc_file)
                
                if result.value:
                    # Убираем HTML теги
                    text = re.sub('<[^<]+?>', '', result.value)
                    text = text.replace('&nbsp;', ' ').replace('&amp;', '&')
                    
                    return {
                        'success': True,
                        'text': text.strip(),
                        'method': 'mammoth',
                        'warnings': result.messages
                    }
                else:
                    return {
                        'success': False,
                        'error': 'DOC файл не содержит извлекаемого текста'
                    }
                    
        except Exception as e:
            # Fallback - предлагаем конвертировать в DOCX
            return {
                'success': False,
                'error': f'Ошибка чтения DOC файла: {str(e)}. Попробуйте сохранить файл в формате DOCX'
            }
    
    def extract_text_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из TXT файла с автоопределением кодировки"""
        try:
            # Определяем кодировку файла
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding']
                
            if not encoding:
                encoding = 'utf-8'  # Fallback
            
            # Читаем файл с определенной кодировкой
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    
                return {
                    'success': True,
                    'text': text.strip(),
                    'method': f'text file ({encoding})'
                }
                
            except UnicodeDecodeError:
                # Пробуем другие популярные кодировки
                encodings_to_try = ['utf-8', 'cp1251', 'cp866', 'iso-8859-1', 'latin-1']
                
                for enc in encodings_to_try:
                    try:
                        with open(file_path, 'r', encoding=enc) as file:
                            text = file.read()
                            
                        return {
                            'success': True,
                            'text': text.strip(),
                            'method': f'text file ({enc})'
                        }
                    except:
                        continue
                
                return {
                    'success': False,
                    'error': 'Не удалось определить кодировку текстового файла'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'Ошибка чтения TXT файла: {str(e)}'
            }

    def extract_text_from_epub(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из EPUB книги"""
        if not HAS_EPUB_SUPPORT:
            return {
                'success': False,
                'error': 'EPUB поддержка недоступна - установите ebooklib'
            }

        try:
            logger.info(f"📚 Начинаю извлечение текста из EPUB: {file_path}")

            book = epub.read_epub(file_path)
            text_content = []
            metadata = {}

            # Извлекаем метаданные книги
            try:
                metadata['title'] = book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else 'Без названия'
                metadata['author'] = book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else 'Неизвестный автор'
                metadata['language'] = book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else 'unknown'
            except Exception as e:
                logger.warning(f"📚 Ошибка извлечения метаданных EPUB: {e}")
                metadata['title'] = 'Без названия'
                metadata['author'] = 'Неизвестный автор'

            # Извлекаем текст из всех документов
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    content = item.get_content()
                    soup = BeautifulSoup(content, 'html.parser')

                    # Удаляем скрипты и стили
                    for script in soup(['script', 'style']):
                        script.decompose()

                    # Извлекаем текст
                    text = soup.get_text(separator=' ', strip=True)
                    if text:
                        text_content.append(text)

            full_text = '\n\n'.join(text_content)

            if not full_text or len(full_text) < 100:
                return {
                    'success': False,
                    'error': 'EPUB файл не содержит текста или текст слишком короткий'
                }

            logger.info(f"📚 EPUB успешно обработан: {len(full_text)} символов")
            logger.info(f"📚 Метаданные: {metadata['title']} - {metadata['author']}")

            return {
                'success': True,
                'text': full_text,
                'method': 'ebooklib',
                'meta': metadata
            }

        except Exception as e:
            logger.error(f"📚 Ошибка чтения EPUB: {e}")
            return {
                'success': False,
                'error': f'Ошибка чтения EPUB: {str(e)}'
            }

    def extract_text_from_fb2(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из FB2 книги"""
        if not HAS_FB2_SUPPORT:
            return {
                'success': False,
                'error': 'FB2 поддержка недоступна - XML парсер недоступен'
            }

        try:
            logger.info(f"📚 Начинаю извлечение текста из FB2: {file_path}")

            # Читаем файл
            with open(file_path, 'rb') as f:
                content = f.read()

            # Определяем кодировку
            detected = chardet.detect(content)
            encoding = detected['encoding'] or 'utf-8'

            # Парсим XML
            tree = ET.parse(file_path)
            root = tree.getroot()

            # FB2 использует namespace
            namespaces = {'fb': 'http://www.gribuser.ru/xml/fictionbook/2.0'}

            # Извлекаем метаданные
            metadata = {}
            try:
                title_info = root.find('.//fb:title-info', namespaces)
                if title_info is not None:
                    book_title = title_info.find('fb:book-title', namespaces)
                    metadata['title'] = book_title.text if book_title is not None else 'Без названия'

                    authors = title_info.findall('.//fb:author', namespaces)
                    author_names = []
                    for author in authors:
                        first_name = author.find('fb:first-name', namespaces)
                        last_name = author.find('fb:last-name', namespaces)
                        if first_name is not None and last_name is not None:
                            author_names.append(f"{first_name.text} {last_name.text}")
                    metadata['author'] = ', '.join(author_names) if author_names else 'Неизвестный автор'
                else:
                    metadata['title'] = 'Без названия'
                    metadata['author'] = 'Неизвестный автор'
            except Exception as e:
                logger.warning(f"📚 Ошибка извлечения метаданных FB2: {e}")
                metadata['title'] = 'Без названия'
                metadata['author'] = 'Неизвестный автор'

            # Извлекаем текст из body
            text_parts = []
            body = root.find('.//fb:body', namespaces)
            if body is not None:
                # Рекурсивно извлекаем весь текст
                for elem in body.iter():
                    if elem.text:
                        text_parts.append(elem.text.strip())
                    if elem.tail:
                        text_parts.append(elem.tail.strip())

            # Объединяем текст
            full_text = '\n'.join([t for t in text_parts if t])

            if not full_text or len(full_text) < 100:
                return {
                    'success': False,
                    'error': 'FB2 файл не содержит текста или текст слишком короткий'
                }

            logger.info(f"📚 FB2 успешно обработан: {len(full_text)} символов")
            logger.info(f"📚 Метаданные: {metadata['title']} - {metadata['author']}")

            return {
                'success': True,
                'text': full_text,
                'method': 'xml.etree',
                'meta': metadata
            }

        except Exception as e:
            logger.error(f"📚 Ошибка чтения FB2: {e}")
            return {
                'success': False,
                'error': f'Ошибка чтения FB2: {str(e)}'
            }

    def extract_text_from_file(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Универсальная функция извлечения текста из файлов"""

        # Нормализуем расширение
        extension = file_extension.lower()

        # Выбираем метод в зависимости от расширения
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.pptx':
            return self.extract_text_from_pptx(file_path)
        elif extension in ('.png', '.jpg', '.jpeg'):
            return self.extract_text_from_image(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif extension == '.doc':
            return self.extract_text_from_doc(file_path)
        elif extension == '.txt':
            return self.extract_text_from_txt(file_path)
        elif extension == '.epub':
            return self.extract_text_from_epub(file_path)
        elif extension == '.fb2':
            return self.extract_text_from_fb2(file_path)
        else:
            return {
                'success': False,
                'error': f'Неподдерживаемый формат файла: {extension}'
            }
    
    def extract_text_from_pptx(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из PPTX презентации"""
        if not HAS_PPTX_SUPPORT:
            return {
                'success': False,
                'error': 'PPTX поддержка недоступна - установите python-pptx'
            }
        
        logger.info(f"📊 Начинаю извлечение текста из PPTX: {file_path}")
        return extract_text_from_pptx(file_path)

    def extract_text_from_image(self, file_path: str) -> Dict[str, Any]:
        """Извлекает текст из изображения с помощью OCR"""
        if not HAS_ENHANCED_PDF_SUPPORT:
            return {
                'success': False,
                'error': 'OCR поддержка недоступна - установите pytesseract и Tesseract'
            }
        
        try:
            import pytesseract
            from PIL import Image
            
            logger.info(f"🖼️ Начинаю OCR изображения: {file_path}")
            
            img = Image.open(file_path)
            ocr_langs = os.getenv("OCR_LANGS", "rus+eng")
            
            txt = pytesseract.image_to_string(
                img, 
                lang=ocr_langs, 
                config="--psm 6 -c tessedit_char_whitelist=АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯабвгдеёжзийклмнопрстуфхцчшщъыьэюяABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?:;()[]{}\"'- \n"
            ).strip()
            
            if not txt or len(txt) < 10:
                return {
                    'success': False,
                    'error': 'Не удалось распознать текст на изображении'
                }
            
            logger.info(f"🖼️ OCR извлек {len(txt)} символов")
            
            return {
                'success': True,
                'text': txt,
                'method': 'image-ocr',
                'meta': {
                    'ocr_language': ocr_langs,
                    'chars_extracted': len(txt)
                }
            }
            
        except Exception as e:
            logger.error(f"🖼️ OCR ошибка: {e}")
            return {
                'success': False,
                'error': f'OCR ошибка: {str(e)}'
            }

    def cleanup_temp_file(self, temp_dir: str):
        """Очищает временные файлы"""
        try:
            import shutil
            shutil.rmtree(temp_dir)
            logger.info(f"Временная директория удалена: {temp_dir}")
        except Exception as e:
            logger.error(f"Ошибка при удалении временной директории {temp_dir}: {e}")